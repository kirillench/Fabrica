import csv
import io
import json
from datetime import date


def to_json(goal: str, hypotheses: list[dict]) -> bytes:
    payload = {"goal": goal, "generated_at": date.today().isoformat(), "hypotheses": hypotheses}
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def to_csv(goal: str, hypotheses: list[dict]) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf, delimiter=";")
    writer.writerow(
        ["Ранг", "Название", "Гипотеза", "Балл", "Новизна", "Реализуемость",
         "Эффект", "Риск", "Риски (описание)", "Критерий успеха", "Источники"]
    )
    for i, h in enumerate(hypotheses, 1):
        sources = ", ".join(s.get("filename", "") for s in h.get("sources", []))
        writer.writerow(
            [i, h.get("title", ""), h.get("statement", ""), h.get("score", ""),
             h.get("novelty", ""), h.get("feasibility", ""), h.get("impact", ""),
             h.get("risk", ""), h.get("risks_text", ""),
             h.get("success_criteria", ""), sources]
        )
    # BOM, чтобы Excel корректно открыл кириллицу
    return ("﻿" + buf.getvalue()).encode("utf-8")


def to_docx(goal: str, hypotheses: list[dict]) -> bytes:
    import docx
    from docx.shared import Pt

    doc = docx.Document()
    doc.add_heading("Фабрика гипотез — отчёт", level=0)
    doc.add_paragraph(f"Дата: {date.today().isoformat()}")
    doc.add_paragraph(f"Целевая проблема: {goal}")

    # сводная таблица имеет смысл только для набора гипотез
    if len(hypotheses) > 1:
        doc.add_heading("Сводное ранжирование", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, name in enumerate(["Ранг", "Гипотеза", "Балл", "Риск"]):
            hdr[i].text = name
        for i, h in enumerate(hypotheses, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = h.get("title", "")
            row[2].text = str(h.get("score", ""))
            row[3].text = str(h.get("risk", ""))

    for i, h in enumerate(hypotheses, 1):
        doc.add_heading(f"{i}. {h.get('title', '')}", level=1)
        doc.add_paragraph(h.get("statement", ""))

        p = doc.add_paragraph()
        p.add_run("Оценки: ").bold = True
        p.add_run(
            f"итог {h.get('score')}/100 · новизна {h.get('novelty')}/10 · "
            f"реализуемость {h.get('feasibility')}/10 · эффект {h.get('impact')}/10 · "
            f"риск {h.get('risk')}/10"
        )

        for label, key in (
            ("Механизм влияния", "mechanism"),
            ("Обоснование", "rationale"),
            ("Риски", "risks_text"),
            ("Критерий успеха", "success_criteria"),
        ):
            if h.get(key):
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(h[key]))

        if h.get("roadmap"):
            p = doc.add_paragraph()
            p.add_run("Дорожная карта проверки:").bold = True
            total_days = 0
            total_cost = 0
            for step in h["roadmap"]:
                if isinstance(step, dict):
                    days = int(step.get("duration_days") or 0)
                    cost = int(step.get("cost") or 0)
                    total_days += days
                    total_cost += cost
                    line = f"{step.get('name', '')} — {days} дн."
                    if cost:
                        line += f", {cost:,} ₽".replace(",", " ")
                    if step.get("resources"):
                        line += f" ({step['resources']})"
                    doc.add_paragraph(line, style="List Number")
                else:
                    doc.add_paragraph(str(step), style="List Number")
            if total_days or total_cost:
                p = doc.add_paragraph()
                p.add_run(
                    f"Итого: {total_days} дн., {total_cost:,} ₽".replace(",", " ")
                ).bold = True

        if h.get("sources"):
            p = doc.add_paragraph()
            p.add_run("Источники:").bold = True
            for s in h["sources"]:
                sp = doc.add_paragraph(style="List Bullet")
                sp.add_run(f"{s.get('filename', '')}: ").bold = True
                run = sp.add_run(f"«{s.get('snippet', '')}…»")
                run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

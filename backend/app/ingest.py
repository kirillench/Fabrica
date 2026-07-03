import io


def extract_text(filename: str, data: bytes) -> str:
    """Извлекает текст из PDF/DOCX/XLSX/CSV/TXT/MD."""
    name = filename.lower()
    if name.endswith(".pdf"):
        return _from_pdf(data)
    if name.endswith(".docx"):
        return _from_docx(data)
    if name.endswith((".xlsx", ".xlsm")):
        return _from_xlsx(data)
    # txt, md, csv и всё остальное читаем как текст
    return data.decode("utf-8", errors="replace")


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n\n".join((page.extract_text() or "") for page in reader.pages)


def _from_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _from_xlsx(data: bytes) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines: list[str] = []
    for ws in wb.worksheets:
        lines.append(f"### Лист: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def chunk_text(text: str, size: int = 1200, overlap: int = 200) -> list[str]:
    """Режет текст на перекрывающиеся чанки, стараясь рвать по границам абзацев."""
    text = text.strip()
    if not text:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + size, len(text))
        if end < len(text):
            # ищем ближайший разрыв абзаца/предложения в хвосте чанка
            for sep in ("\n\n", "\n", ". "):
                cut = text.rfind(sep, start + size // 2, end)
                if cut != -1:
                    end = cut + len(sep)
                    break
        piece = text[start:end].strip()
        if piece:
            chunks.append(piece)
        if end >= len(text):
            break
        start = max(end - overlap, start + 1)
    return chunks
